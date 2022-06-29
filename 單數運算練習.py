from pathlib import Path
wdir = Path(__file__).parent
wordexe = 'C:\Program Files\Microsoft Office\Office14\WINWORD.EXE'

def 單數運算評量(運算='X', 列印=None, 固定左數=None):
    運算名稱 = {'X':'乘法', '+':'加法'}[運算]
    if 固定左數:
        固定左數 = int(固定左數)
        table = [(固定左數, y) for y in range(1,10)]
        table.extend(table*8)
    else:
        table = [(x,y) for x in range(1,10) for y in range(1,10)]

    def 顯示或列印():
        from random import shuffle
        shuffle(table)
        from docx import Document
        f = wdir / '模版.docx'
        doc = Document(f)
        title = f"九九{運算名稱}評量"
        if 固定左數:
            title += f'，固定左數為{固定左數}' 
            if 運算名稱 == '加法':
                title += f'，口訣{固定左數}缺{10-固定左數}' 
        doc.paragraphs[0].text = title
        doc.paragraphs[0].style = doc.styles['Heading 1']
        tab = doc.add_table(rows=9, cols=9)
        for i, c in enumerate(table):
            n1, n2 = c
            c = tab.cell(i%9, i//9)
            if 運算=='X':
                c.text = (f'{n1}X{n2}=(    )')
            elif 運算=='+': 
                c.text = (f'{n1}+{n2}=(    )')
            else:
                raise ValueError('不支援【{運算}】運算題目！')
        doc.add_paragraph()
        doc.add_paragraph()
        result = '答錯題數：____題；答題時間：____分鐘____秒；評量日期：__年__月__日；姓名：________'
        doc.add_paragraph(result)
        doc.add_page_break()
        doc.add_paragraph('答案頁', style='Heading 1')
        tab = doc.add_table(rows=9, cols=9)
        for i, c in enumerate(table):
            n1, n2 = c
            c = tab.cell(i%9, i//9)
            if 運算=='X':
                c.text = (f'{n1}X{n2}=({n1*n2})')
            elif 運算=='+': 
                # breakpoint()
                c.text = (f'{n1}+{n2}=({n1+n2})')
        fn = wdir / f'九九{運算名稱}評量.docx'
        doc.save(fn)
        from os import system
        cmd = f'start {fn}'
        if 列印:
            cmd = f'"{wordexe}" {fn} /mFilePrintDefault /mFileCloseOrExit /q /n'
        system(cmd)
        
    if 列印:
        copy = int(列印) 
        for i in range(0, copy):
            顯示或列印()
    else:
        顯示或列印()

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("operator", help="運算類型")
    parser.add_argument("--print", help="列印份數")
    parser.add_argument("--fix", help="固定左數")
    args = parser.parse_args()
    單數運算評量(args.operator, 列印=args.print, 固定左數=args.fix)
